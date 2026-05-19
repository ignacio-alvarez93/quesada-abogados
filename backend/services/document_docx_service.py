import re
import json
from datetime import datetime
from pathlib import Path

from docx import Document

from backend.services import document_generation_service
from backend.services import document_template_service


BASE_DIR = Path(__file__).resolve().parents[2]
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _normalize_code(value):
    return str(value or "").strip().upper().replace(" ", "_")


def _safe_filename(value):
    raw = _normalize_code(value)
    safe = []
    for char in raw:
        if char.isalnum() or char in ("_", "-"):
            safe.append(char)
        else:
            safe.append("_")
    return "".join(safe).strip("_") or "DOCUMENTO"


def _resolve_project_path(value):
    raw = str(value or "").strip().replace("\\", "/").strip()
    if not raw:
        raise ValueError("La plantilla documental no tiene template_path configurado")

    path = Path(raw)
    if path.is_absolute():
        return path

    return BASE_DIR / path


def _relative_path(path):
    try:
        return str(Path(path).resolve().relative_to(BASE_DIR)).replace("\\", "/")
    except Exception:
        return str(path).replace("\\", "/")


def _stringify(value):
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _payload_lookup(payload, key):
    """
    Busca una clave de placeholder en el payload.

    Soporta:
    - coincidencia literal: {{tipo tramite}}
    - coincidencia normalizada con espacios/guiones: {{tipo_tramite}}
    - rutas simples si el payload contiene dicts: {{cliente.nombre}}
    """
    raw_key = str(key or "").strip()
    if not raw_key:
        return ""

    if raw_key in payload:
        return _stringify(payload.get(raw_key))

    normalized_key = raw_key.replace("_", " ").strip().lower()
    for payload_key, payload_value in (payload or {}).items():
        if str(payload_key).replace("_", " ").strip().lower() == normalized_key:
            return _stringify(payload_value)

    if "." in raw_key:
        current = payload
        for part in raw_key.split("."):
            if isinstance(current, dict) and part in current:
                current = current[part]
            else:
                return ""
        return _stringify(current)

    return ""


def _replace_placeholders_in_text(text, payload):
    if not text:
        return text

    def replace(match):
        key = match.group(1)
        return _payload_lookup(payload, key)

    return PLACEHOLDER_RE.sub(replace, text)


def _set_paragraph_text_preserving_first_run(paragraph, new_text):
    """
    Sustitución básica.

    python-docx puede dividir un placeholder entre varios runs. Para garantizar
    reemplazo en DT-5A, reconstruimos el párrafo con el estilo del primer run.
    En una fase posterior se podrá crear un motor que preserve formato run a run.
    """
    if paragraph.text == new_text:
        return

    first_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    run = paragraph.add_run(new_text)

    if first_run is not None:
        try:
            run.bold = first_run.bold
            run.italic = first_run.italic
            run.underline = first_run.underline
            run.font.name = first_run.font.name
            run.font.size = first_run.font.size
        except Exception:
            pass


def _replace_in_paragraph(paragraph, payload):
    original = paragraph.text
    replaced = _replace_placeholders_in_text(original, payload)
    if replaced != original:
        _set_paragraph_text_preserving_first_run(paragraph, replaced)


def _replace_in_table(table, payload):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, payload)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, payload)


def _replace_docx_placeholders(document, payload):
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, payload)

    for table in document.tables:
        _replace_in_table(table, payload)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, payload)
        for table in section.header.tables:
            _replace_in_table(table, payload)

        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, payload)
        for table in section.footer.tables:
            _replace_in_table(table, payload)


def _collect_unresolved_placeholders(document):
    found = set()

    def collect_from_text(text):
        for match in PLACEHOLDER_RE.findall(text or ""):
            found.add(match.strip())

    for paragraph in document.paragraphs:
        collect_from_text(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_from_text(paragraph.text)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            collect_from_text(paragraph.text)
        for paragraph in section.footer.paragraphs:
            collect_from_text(paragraph.text)

    return sorted(found)


def _build_output_docx_path(export_data):
    document_template = export_data.get("document_template") or {}
    expediente = export_data.get("expediente") or {}
    output = export_data.get("output") or {}

    output_dir = BASE_DIR / str(output.get("directory") or "").replace("\\", "/")
    if not str(output.get("directory") or "").strip():
        raise ValueError("El export de payload no contiene directorio de salida")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    template_code = _safe_filename(document_template.get("codigo"))
    expediente_part = _safe_filename(expediente.get("numero_expediente")) if expediente else "GENERAL"

    return output_dir / f"{template_code}_{expediente_part}_{timestamp}.docx"


def generate_docx_from_template(document_template_id, expediente_id=None, auto_build_snapshot=True):
    """
    Genera un DOCX desde una plantilla documental registrada.

    Flujo:
    document_template -> mapper_destino -> payload JSON -> DOCX final

    No toca Box.
    No toca Mercurio.
    No modifica expedientes.
    """
    template = document_template_service.get_document_template(document_template_id)
    if not template:
        raise ValueError(f"No existe document_template_id={document_template_id}")

    if not int(template.get("activo") or 0):
        raise ValueError("La plantilla documental está inactiva.")

    if str(template.get("template_type") or "").strip().lower() != "docx":
        raise ValueError("Esta plantilla documental no es de tipo docx.")

    template_path = _resolve_project_path(template.get("template_path"))
    if not template_path.exists():
        raise FileNotFoundError(f"No existe el DOCX de plantilla: {_relative_path(template_path)}")

    export_data = document_generation_service.export_document_payload(
        document_template_id,
        expediente_id=expediente_id,
        auto_build_snapshot=auto_build_snapshot,
    )

    payload = export_data.get("payload") or {}
    document = Document(str(template_path))
    _replace_docx_placeholders(document, payload)

    unresolved = _collect_unresolved_placeholders(document)
    output_docx_path = _build_output_docx_path(export_data)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx_path))

    export_data["generation_type"] = "docx"
    export_data["docx"] = {
        "template_path": _relative_path(template_path),
        "docx_path": _relative_path(output_docx_path),
        "unresolved_placeholders": unresolved,
        "unresolved_count": len(unresolved),
    }
    export_data["output"]["docx_path"] = _relative_path(output_docx_path)
    export_data["output"]["format"] = "docx"

    return export_data


def generate_docx_from_template_code(codigo, expediente_id=None, auto_build_snapshot=True):
    template = document_template_service.get_document_template_by_code(codigo, active_only=True)
    if not template:
        raise ValueError(f"No existe plantilla documental activa con código={codigo}")

    return generate_docx_from_template(
        template["id"],
        expediente_id=expediente_id,
        auto_build_snapshot=auto_build_snapshot,
    )
